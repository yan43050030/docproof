"""Read, proofread, and write Word (.docx) documents with format preservation.

Text is collected in reading order from the document body (including nested
tables) as well as every section's headers and footers, so proofreading covers
the places real documents actually hide typos. Corrections are applied at the
run level via :mod:`docproof.document.run_ops`, which preserves character
formatting and can emit genuine Word tracked changes.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Iterator

if TYPE_CHECKING:
    from docx.document import Document as DocumentType

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from docproof.document.position_mapper import PositionMapper, paragraph_text
from docproof.document import run_ops

# Footnote entries that carry no real text (the separator rules).
_SKIP_NOTE_TYPES = {"separator", "continuationSeparator"}


@dataclass
class CorrectionResult:
    """Result of proofreading a document."""

    original_text: str
    corrected_text: str
    errors: list = field(default_factory=list)  # list of ErrorItem
    paragraph_errors: dict[int, list] = field(default_factory=dict)
    # paragraph_index -> list of ErrorItem


def _iter_block_paragraphs(container) -> Iterator[Paragraph]:
    """Yield every paragraph in a block container in reading order.

    Recurses through tables (and nested tables) so table-cell text is included.
    Falls back to ``.paragraphs`` + ``.tables`` for containers (headers/footers)
    that don't implement ``iter_inner_content``.
    """
    if hasattr(container, "iter_inner_content"):
        for item in container.iter_inner_content():
            if isinstance(item, Paragraph):
                yield item
            elif isinstance(item, Table):
                yield from _iter_table_paragraphs(item)
        return
    for para in getattr(container, "paragraphs", []):
        yield para
    for table in getattr(container, "tables", []):
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_block_paragraphs(cell)


def _iter_textbox_paragraphs(element) -> Iterator[Paragraph]:
    """Yield paragraphs inside every text box (``w:txbxContent``) under element.

    Text box content lives in the main document / header / footer element trees,
    so edits to these paragraphs persist through a normal save.
    """
    if element is None:
        return
    for txbx in element.iter(qn("w:txbxContent")):
        for p in txbx.findall(qn("w:p")):
            yield Paragraph(p, None)


class DocxHandler:
    """Handles .docx document reading, correction application, and writing."""

    def __init__(self, filepath: str):
        self.filepath = filepath
        self._doc: DocumentType | None = None
        self._mapper = PositionMapper()
        self._paragraphs: list[Paragraph] = []
        self._result: CorrectionResult | None = None
        # (part, parsed_root) for footnote/endnote parts that need their blob
        # rewritten on save (python-docx keeps them as opaque blobs).
        self._note_parts: list = []

    # ---- properties ----

    @property
    def document(self):
        return self._doc

    @property
    def paragraphs(self) -> list:
        return list(self._paragraphs)

    @property
    def mapper(self) -> PositionMapper:
        return self._mapper

    @property
    def result(self) -> CorrectionResult | None:
        return self._result

    # ---- loading ----

    def load(self) -> None:
        """Open and parse the document, collecting all proofread-able text."""
        self._doc = Document(self.filepath)
        self._paragraphs = self._collect_paragraphs()
        self._mapper.build(self._paragraphs)

    def _collect_paragraphs(self) -> list[Paragraph]:
        """Collect proofread-able paragraphs from every part of the document.

        Order: body (with tables), body text boxes, then each section's headers
        and footers (with their text boxes), then footnotes and endnotes.
        """
        self._note_parts = []
        paras: list[Paragraph] = []
        paras.extend(_iter_block_paragraphs(self._doc))
        paras.extend(_iter_textbox_paragraphs(self._doc.element.body))
        for section in self._doc.sections:
            for hf in (
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer,
            ):
                if hf is None or getattr(hf, "is_linked_to_previous", False):
                    continue
                paras.extend(_iter_block_paragraphs(hf))
                paras.extend(_iter_textbox_paragraphs(getattr(hf, "_element", None)))
        paras.extend(self._collect_notes())
        return paras

    def _collect_notes(self) -> list[Paragraph]:
        """Collect footnote/endnote paragraphs, registering their parts so edits
        can be written back on save."""
        paras: list[Paragraph] = []
        try:
            rels = list(self._doc.part.rels.values())
        except Exception:
            return paras
        for rel in rels:
            reltype = getattr(rel, "reltype", "")
            if not (reltype.endswith("/footnotes") or reltype.endswith("/endnotes")):
                continue
            try:
                root = parse_xml(rel.target_part.blob)
            except Exception:
                continue
            for note in root:
                if note.get(qn("w:type")) in _SKIP_NOTE_TYPES:
                    continue
                for p in note.findall(qn("w:p")):
                    paras.append(Paragraph(p, None))
            self._note_parts.append((rel.target_part, root))
        return paras

    def get_full_text(self) -> str:
        """Full proofread text: paragraph run-texts joined by newlines."""
        if not self._doc:
            raise RuntimeError("Document not loaded. Call load() first.")
        return "\n".join(self._mapper.paragraph_texts)

    # ---- correction application ----

    def _group_errors(self, errors: list) -> dict[int, list]:
        """Group errors by paragraph index, sorted descending within a paragraph.

        Descending order keeps earlier offsets valid while later spans are edited
        (run splitting/replacement only affects text at or after the edit point).
        """
        grouped: dict[int, list] = {}
        for err in errors:
            span = self._mapper.char_to_span(err.start)
            if span is None:
                continue
            grouped.setdefault(span.paragraph_index, []).append(err)
        for p_idx in grouped:
            grouped[p_idx].sort(key=lambda e: e.start, reverse=True)
        return grouped

    def apply_corrections(
        self, errors: list, markup: bool = True, track_changes: bool = False
    ) -> CorrectionResult:
        """Apply corrections to the document.

        Args:
            errors: ErrorItem objects from the proofreading engine.
            markup: If True (and not track_changes), add visible colored markup.
                    If False, replace the error text cleanly.
            track_changes: If True, emit genuine Word tracked changes that can be
                    accepted/rejected in Word or WPS. Takes precedence over markup.
        """
        if not self._doc:
            raise RuntimeError("Document not loaded. Call load() first.")

        original = self.get_full_text()
        grouped = self._group_errors(errors)

        for p_idx, p_errors in grouped.items():
            para = self._mapper.get_paragraph(p_idx)
            para_start, _ = self._mapper.get_paragraph_range(p_idx)
            for err in p_errors:
                local_start = err.start - para_start
                local_end = err.end - para_start
                if track_changes:
                    run_ops.revise_span(para, local_start, local_end,
                                        err.error, err.correct)
                elif markup:
                    run_ops.markup_span(para, local_start, local_end,
                                        err.error, err.correct)
                else:
                    run_ops.replace_span(para, local_start, local_end, err.correct)

        # Rebuild the mapper so corrected_text and any subsequent edits are valid.
        self._mapper.build(self._paragraphs)
        result = CorrectionResult(
            original_text=original,
            corrected_text=self.get_full_text(),
            errors=errors,
            paragraph_errors=grouped,
        )
        self._result = result
        return result

    # ---- saving ----

    def save(self, filepath: str) -> None:
        """Save the document to a new file.

        Footnote/endnote parts are opaque blobs in python-docx, so their edited
        element trees are serialised back into the parts before saving.
        """
        if not self._doc:
            raise RuntimeError("Document not loaded.")
        for part, root in self._note_parts:
            try:
                part._blob = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
            except Exception:
                pass
        self._doc.save(filepath)

    def replace_full_text(self, new_text: str) -> None:
        """Apply hand-edited full text back to the document (edit mode).

        When the paragraph structure is unchanged (the common case — the user
        edited words, not line breaks), only paragraphs whose text actually
        changed are rewritten, so every untouched paragraph keeps its full
        formatting and every table/text box/header stays intact.

        If the number of lines changed, falls back to a plain body rewrite.
        """
        if not self._doc:
            raise RuntimeError("Document not loaded.")

        new_lines = new_text.split("\n")
        paras = self._paragraphs

        if len(new_lines) == len(paras):
            for para, line in zip(paras, new_lines):
                old = paragraph_text(para)
                if old != line:
                    self._set_paragraph_text(para, old, line)
        else:
            self._rewrite_body(new_lines)

        self._paragraphs = self._collect_paragraphs()
        self._mapper.build(self._paragraphs)

    @staticmethod
    def _set_paragraph_text(para, old: str, line: str) -> None:
        """Replace a paragraph's whole text, keeping the first run's format."""
        if not para.runs:
            para.add_run(line)
        else:
            run_ops.replace_span(para, 0, len(old), line)

    def _rewrite_body(self, new_lines: list[str]) -> None:
        """Fallback plain-text rewrite over body paragraphs when the line count
        changed (structure edits can't be mapped to the structured layout)."""
        body_paras = list(self._doc.paragraphs)
        for i, line in enumerate(new_lines):
            if i < len(body_paras):
                para = body_paras[i]
                if para.runs:
                    para.runs[0].text = line
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(line)
            else:
                self._doc.add_paragraph(line)
        for para in body_paras[len(new_lines):]:
            for run in para.runs:
                run.text = ""
