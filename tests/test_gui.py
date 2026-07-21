"""Headless GUI tests (offscreen Qt) for the main window wiring."""

import os
import tempfile

import pytest

pytest.importorskip("PySide6")

from docx import Document  # noqa: E402

from docproof.engine.engine_manager import EngineManager  # noqa: E402
from docproof.engine.rule_engine import RuleEngine  # noqa: E402

pytestmark = pytest.mark.gui


def _engine():
    em = EngineManager()
    em._engine = RuleEngine()
    em._engine.load()
    return em


def _docx(paragraphs) -> str:
    d = Document()
    for p in paragraphs:
        d.add_paragraph(p)
    path = tempfile.mktemp(suffix=".docx")
    d.save(path)
    return path


@pytest.fixture
def window(qapp):
    from docproof.ui.main_window import MainWindow
    w = MainWindow(_engine())
    yield w
    if w._warmup_thread is not None and w._warmup_thread.isRunning():
        w._warmup_thread.wait(3000)


def _proofread(window):
    text = window._docx_handler.get_full_text()
    window._proofread_text = text
    errs = window._engine_manager.proofread(text)
    window._on_proofread_done(errs)
    return errs


class TestMainWindowFlow:
    def test_load_and_proofread(self, window):
        path = _docx(["你好,世界", "测试 空格,标点"])
        try:
            window._load_document(path)
            errs = _proofread(window)
            assert errs
            assert window._error_list.list_widget.count() == len(errs)
        finally:
            os.remove(path)

    def test_click_selects_list_row(self, window):
        path = _docx(["你好,世界"])
        try:
            window._load_document(path)
            errs = _proofread(window)
            # Emitting error_clicked should select the matching list row.
            window._correction_view.error_clicked.emit(0)
            sel = window._error_list.list_widget.currentItem()
            assert sel is not None
            assert sel.data(256) == 0  # Qt.UserRole
        finally:
            os.remove(path)

    def test_accept_then_undo(self, window):
        path = _docx(["你好,世界"])
        try:
            window._load_document(path)
            _proofread(window)
            window._error_list._accept_by_index(0)
            assert 0 in window._error_list.accepted_indices
            action, idx = window._error_list.undo_last()
            assert action == "accept" and idx == 0
            assert 0 not in window._error_list.accepted_indices
        finally:
            os.remove(path)

    def test_streaming_partial(self, window):
        path = _docx(["你好,世界", "测试 空格"])
        try:
            window._load_document(path)
            text = window._docx_handler.get_full_text()
            window._proofread_text = text
            errs = window._engine_manager.proofread(text)
            window._on_proofread_partial(errs[:1])
            window._on_proofread_partial(errs)
            assert window._correction_view.error_count == len(errs)
        finally:
            os.remove(path)

    def test_status_color_resets(self, window):
        window._set_status("ok", kind="ok")
        assert window._status_text.styleSheet() != ""
        window._set_status("plain")
        assert window._status_text.styleSheet() == ""

    def test_export_clean_idempotent(self, window):
        path = _docx(["你好,世界"])
        try:
            window._load_document(path)
            _proofread(window)
            window._error_list._accept_all()
            handler = window._fresh_handler()
            handler.apply_corrections(
                window._error_list.get_accepted_errors(), markup=False)
            out = tempfile.mktemp(suffix=".docx")
            handler.save(out)
            assert Document(out).paragraphs[0].text == "你好，世界"
            os.remove(out)
        finally:
            os.remove(path)


class TestTheme:
    def test_theme_switch_persists(self, window):
        window._set_theme("dark")
        assert window._settings.get("theme") == "dark"
        window._set_theme("light")
        assert window._settings.get("theme") == "light"
