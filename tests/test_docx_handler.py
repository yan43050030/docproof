"""Tests for DocxHandler: table/header collection and correction application."""

import os
import tempfile

from docx import Document

from docproof.document.docx_handler import DocxHandler
from docproof.engine.base_engine import ErrorItem


def _build_doc() -> str:
    doc = Document()
    doc.add_paragraph("正文里有错字。")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "表格错字一"
    t.rows[0].cells[1].text = "单元格文本"
    doc.sections[0].header.paragraphs[0].text = "页眉错字"
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path


class TestCollection:
    def test_collects_body_table_and_header(self):
        path = _build_doc()
        try:
            h = DocxHandler(path)
            h.load()
            full = h.get_full_text()
            assert "正文里有错字" in full
            assert "表格错字一" in full  # table cell
            assert "页眉错字" in full     # header
        finally:
            os.remove(path)


class TestApplyCorrections:
    def test_direct_replace_in_table(self):
        path = _build_doc()
        try:
            h = DocxHandler(path)
            h.load()
            full = h.get_full_text()
            idx = full.index("表格错字一") + 2
            err = ErrorItem("错字", "错别字", idx, idx + 2)
            h.apply_corrections([err], markup=False)
            out = tempfile.mktemp(suffix=".docx")
            h.save(out)
            assert Document(out).tables[0].rows[0].cells[0].text == "表格错别字一"
            os.remove(out)
        finally:
            os.remove(path)

    def test_track_changes_in_header(self):
        path = _build_doc()
        try:
            h = DocxHandler(path)
            h.load()
            full = h.get_full_text()
            idx = full.index("页眉错字") + 2
            err = ErrorItem("错字", "错别字", idx, idx + 2)
            h.apply_corrections([err], track_changes=True)
            out = tempfile.mktemp(suffix=".docx")
            h.save(out)
            hxml = Document(out).sections[0].header.paragraphs[0]._p.xml
            assert "w:ins" in hxml and "w:del" in hxml
            os.remove(out)
        finally:
            os.remove(path)

    def test_last_error_not_dropped(self):
        """Regression: corrections on the final paragraph must be applied."""
        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("第二段有错")
        path = tempfile.mktemp(suffix=".docx")
        doc.save(path)
        try:
            h = DocxHandler(path)
            h.load()
            full = h.get_full_text()
            idx = full.index("有错")
            err = ErrorItem("有错", "有误", idx, idx + 2)
            h.apply_corrections([err], markup=False)
            out = tempfile.mktemp(suffix=".docx")
            h.save(out)
            assert Document(out).paragraphs[1].text == "第二段有误"
            os.remove(out)
        finally:
            os.remove(path)
