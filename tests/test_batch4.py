"""Tests for batch-4 features: text boxes, footnotes, diff export, CLI, download."""

import os
import tempfile
import zipfile

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from docproof.document.docx_handler import DocxHandler
from docproof.engine.base_engine import ErrorItem


def _textbox_docx() -> str:
    d = Document()
    p = d.add_paragraph("正文段落")
    run = p.add_run()
    txbx = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml"><v:shape><v:textbox><w:txbxContent>'
        '<w:p><w:r><w:t>文本框错字</w:t></w:r></w:p>'
        '</w:txbxContent></v:textbox></v:shape></w:pict>'
    )
    run._r.append(parse_xml(txbx))
    path = tempfile.mktemp(suffix=".docx")
    d.save(path)
    return path


def _footnote_docx() -> str:
    d = Document()
    d.add_paragraph("正文")
    base = tempfile.mktemp(suffix=".docx")
    d.save(base)
    work = tempfile.mkdtemp()
    with zipfile.ZipFile(base) as z:
        z.extractall(work)
    foot = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:id="1"><w:p><w:r><w:t>脚注错字</w:t></w:r></w:p></w:footnote>'
            '</w:footnotes>')
    with open(os.path.join(work, "word", "footnotes.xml"), "w", encoding="utf-8") as f:
        f.write(foot)
    ct = os.path.join(work, "[Content_Types].xml")
    with open(ct, encoding="utf-8") as f:
        s = f.read()
    s = s.replace("</Types>",
                  '<Override PartName="/word/footnotes.xml" ContentType="application/'
                  'vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/></Types>')
    with open(ct, "w", encoding="utf-8") as f:
        f.write(s)
    rels = os.path.join(work, "word", "_rels", "document.xml.rels")
    with open(rels, encoding="utf-8") as f:
        r = f.read()
    r = r.replace("</Relationships>",
                  '<Relationship Id="rIdFoot" Type="http://schemas.openxmlformats.org/'
                  'officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/></Relationships>')
    with open(rels, "w", encoding="utf-8") as f:
        f.write(r)
    out = tempfile.mktemp(suffix=".docx")
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _d, files in os.walk(work):
            for fn in files:
                fp = os.path.join(root, fn)
                z.write(fp, os.path.relpath(fp, work))
    return out


class TestTextBox:
    def test_collected_and_corrected(self):
        path = _textbox_docx()
        try:
            h = DocxHandler(path)
            h.load()
            full = h.get_full_text()
            assert "文本框错字" in full
            idx = full.index("错字", full.index("文本框"))
            h.apply_corrections([ErrorItem("错字", "错别字", idx, idx + 2)],
                                markup=False)
            out = tempfile.mktemp(suffix=".docx")
            h.save(out)
            d2 = Document(out)
            texts = [Paragraph(p, None).text
                     for tb in d2.element.body.iter(qn("w:txbxContent"))
                     for p in tb.findall(qn("w:p"))]
            assert texts == ["文本框错别字"]
        finally:
            os.remove(path)


class TestFootnotes:
    def test_collected_and_corrected(self):
        path = _footnote_docx()
        try:
            h = DocxHandler(path)
            h.load()
            full = h.get_full_text()
            assert "脚注错字" in full
            idx = full.index("错字", full.index("脚注"))
            h.apply_corrections([ErrorItem("错字", "错别字", idx, idx + 2)],
                                markup=False)
            out = tempfile.mktemp(suffix=".docx")
            h.save(out)
            d2 = Document(out)
            fp = [rel.target_part for rel in d2.part.rels.values()
                  if rel.reltype.endswith("/footnotes")][0]
            root = parse_xml(fp.blob)
            texts = [Paragraph(p, None).text
                     for note in root for p in note.findall(qn("w:p"))
                     if note.get(qn("w:type")) not in ("separator", "continuationSeparator")]
            assert texts == ["脚注错别字"]
        finally:
            os.remove(path)


class TestDiffExport:
    def test_only_changed_paragraph_touched(self):
        d = Document()
        p0 = d.add_paragraph()
        p0.add_run("第一段").bold = True
        p1 = d.add_paragraph("第二段有错字")
        path = tempfile.mktemp(suffix=".docx")
        d.save(path)
        try:
            h = DocxHandler(path)
            h.load()
            full = h.get_full_text()
            # Change only the second paragraph.
            edited = full.replace("有错字", "有错别字")
            h.replace_full_text(edited)
            out = tempfile.mktemp(suffix=".docx")
            h.save(out)
            d2 = Document(out)
            assert d2.paragraphs[0].text == "第一段"
            # Unchanged paragraph kept its bold run formatting.
            assert d2.paragraphs[0].runs[0].bold is True
            assert d2.paragraphs[1].text == "第二段有错别字"
        finally:
            os.remove(path)


class TestCLI:
    def test_batch_clean(self):
        d = Document()
        d.add_paragraph("你好,世界")  # ascii comma -> rule engine fix
        src_dir = tempfile.mkdtemp()
        d.save(os.path.join(src_dir, "a.docx"))
        out_dir = tempfile.mkdtemp()

        from docproof.cli import run
        # Use rule engine only path by loading no model? CLI needs a model;
        # but rule + fix run regardless. Provide a fake via monkeypatch-free path:
        # auto_load may fail without a model, so we call the internal file proc.
        from docproof.engine.engine_manager import EngineManager
        from docproof.engine.rule_engine import RuleEngine
        from docproof.engine.user_dict import UserDict
        from docproof import cli
        em = EngineManager()
        em._engine = RuleEngine()
        em._engine.load()
        n = cli._process_file(
            os.path.join(src_dir, "a.docx"),
            os.path.join(out_dir, "a-校对.docx"),
            "clean", em, UserDict(dict_path=tempfile.mktemp(suffix=".txt")))
        assert n >= 1
        assert Document(os.path.join(out_dir, "a-校对.docx")).paragraphs[0].text == "你好，世界"


class TestDownloader:
    def test_download_atomic(self):
        import http.server
        import socketserver
        import threading

        payload = b"MODELDATA" * 1000
        handler_dir = tempfile.mkdtemp()
        with open(os.path.join(handler_dir, "m.bin"), "wb") as f:
            f.write(payload)

        class Handler(http.server.SimpleHTTPRequestHandler):
            def __init__(self, *a, **k):
                super().__init__(*a, directory=handler_dir, **k)

            def log_message(self, *a):
                pass

        with socketserver.TCPServer(("127.0.0.1", 0), Handler) as httpd:
            port = httpd.server_address[1]
            t = threading.Thread(target=httpd.serve_forever, daemon=True)
            t.start()
            try:
                from docproof.downloader import download
                dest = os.path.join(tempfile.mkdtemp(), "sub", "m.bin")
                seen = []
                download(f"http://127.0.0.1:{port}/m.bin", dest,
                         progress=lambda d, tot: seen.append((d, tot)))
                assert os.path.exists(dest)
                with open(dest, "rb") as f:
                    assert f.read() == payload
                assert not os.path.exists(dest + ".part")
            finally:
                httpd.shutdown()

    def test_download_cancel_leaves_no_file(self):
        import http.server
        import socketserver
        import threading

        handler_dir = tempfile.mkdtemp()
        with open(os.path.join(handler_dir, "big.bin"), "wb") as f:
            f.write(b"x" * 500000)

        class Handler(http.server.SimpleHTTPRequestHandler):
            def __init__(self, *a, **k):
                super().__init__(*a, directory=handler_dir, **k)

            def log_message(self, *a):
                pass

        with socketserver.TCPServer(("127.0.0.1", 0), Handler) as httpd:
            port = httpd.server_address[1]
            t = threading.Thread(target=httpd.serve_forever, daemon=True)
            t.start()
            try:
                from docproof.downloader import download
                dest = os.path.join(tempfile.mkdtemp(), "big.bin")
                try:
                    download(f"http://127.0.0.1:{port}/big.bin", dest,
                             should_stop=lambda: True)
                except InterruptedError:
                    pass
                assert not os.path.exists(dest)
                assert not os.path.exists(dest + ".part")
            finally:
                httpd.shutdown()
