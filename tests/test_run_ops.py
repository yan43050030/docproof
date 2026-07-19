"""Tests for run-level document editing (docproof.document.run_ops)."""

import io

from docx import Document

from docproof.document import run_ops
from docproof.document.position_mapper import paragraph_text


def _roundtrip(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


def _multi_run_para():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("这是")
    r1.bold = True
    r2 = p.add_run("错字")
    r2.italic = True
    p.add_run("测试")
    return doc, p


class TestReplaceSpan:
    def test_replace_middle(self):
        doc, p = _multi_run_para()
        assert run_ops.replace_span(p, 2, 4, "错别字")
        assert paragraph_text(p) == "这是错别字测试"

    def test_replace_preserves_surrounding_format(self):
        doc, p = _multi_run_para()
        run_ops.replace_span(p, 2, 4, "X")
        # First run stays bold, its text intact.
        assert p.runs[0].text == "这是"
        assert p.runs[0].bold is True

    def test_replace_at_start_and_end(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("abcdef")
        run_ops.replace_span(p, 0, 2, "XY")
        assert paragraph_text(p) == "XYcdef"
        run_ops.replace_span(p, 4, 6, "ZZ")
        assert paragraph_text(p) == "XYcdZZ"

    def test_roundtrip_saves(self):
        doc, p = _multi_run_para()
        run_ops.replace_span(p, 2, 4, "错别字")
        d2 = _roundtrip(doc)
        assert d2.paragraphs[0].text == "这是错别字测试"


class TestMarkupSpan:
    def test_markup_inserts_correction(self):
        doc, p = _multi_run_para()
        assert run_ops.markup_span(p, 2, 4, "错字", "错别字")
        txt = paragraph_text(p)
        assert "错字" in txt and "→" in txt and "错别字" in txt

    def test_markup_strikes_error(self):
        doc, p = _multi_run_para()
        run_ops.markup_span(p, 2, 4, "错字", "错别字")
        assert any(r.font.strike for r in p.runs)


class TestReviseSpan:
    def test_revise_produces_tracked_changes(self):
        doc, p = _multi_run_para()
        assert run_ops.revise_span(p, 2, 4, "错字", "错别字",
                                   when="2020-01-01T00:00:00Z")
        xml = _roundtrip(doc).paragraphs[0]._p.xml
        assert "w:ins" in xml
        assert "w:del" in xml
        assert "delText" in xml
        # Deleted and inserted text are both present in the XML.
        assert "错字" in xml and "错别字" in xml

    def test_revise_keeps_visible_text_outside_span(self):
        doc, p = _multi_run_para()
        run_ops.revise_span(p, 2, 4, "错字", "错别字")
        # python-docx para.text ignores ins/del runs -> only untouched text.
        assert _roundtrip(doc).paragraphs[0].text == "这是测试"


class TestIsolateRuns:
    def test_empty_span(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("abc")
        assert run_ops.isolate_runs(p, 1, 1) == []
